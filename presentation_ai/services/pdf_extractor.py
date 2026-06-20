import io
import re
import fitz   # PyMuPDF
from docx import Document

MAX_SOURCE_CHARS = 12000


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a PDF, DOCX, or TXT file given its raw bytes and filename."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext == "docx":
        return _extract_docx(file_bytes)
    elif ext == "doc":
        return _extract_legacy_doc(file_bytes)
    elif ext == "txt":
        return _extract_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: .pdf, .doc, .docx, .txt")


def _extract_pdf(pdf_bytes: bytes) -> str:
    """Extract all text from a PDF."""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages.append(text.strip())
    doc.close()
    full_text = "\n\n".join(pages)
    return full_text[:MAX_SOURCE_CHARS]


def _extract_docx(docx_bytes: bytes) -> str:
    """Extract all text from a DOCX file, including tables."""
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = []
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
                    
    full_text = "\n\n".join(paragraphs)
    return full_text[:MAX_SOURCE_CHARS]


def _extract_legacy_doc(doc_bytes: bytes) -> str:
    """Best-effort text extraction for old binary .doc files."""
    candidates = []
    for encoding in ("utf-16le", "latin-1"):
        decoded = doc_bytes.decode(encoding, errors="ignore")
        runs = re.findall(r"[A-Za-z0-9][\w\s.,;:!?()/%+\-]{4,}", decoded)
        cleaned = [_normalize_space(run) for run in runs]
        text = "\n".join(run for run in cleaned if len(run) >= 5)
        if text:
            candidates.append(text)

    if not candidates:
        raise ValueError(
            "Unable to extract text from this .doc file. Please upload a .docx, "
            "text-based PDF, or TXT file."
        )

    best_text = max(candidates, key=len)
    return best_text[:MAX_SOURCE_CHARS]


def _extract_txt(txt_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    full_text = txt_bytes.decode("utf-8", errors="ignore")
    return full_text[:MAX_SOURCE_CHARS]


def _normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()
